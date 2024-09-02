import base64
import zipfile
from io import BytesIO
from functools import partial
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from docx import Document
from docx.shared import Mm
from bs4 import BeautifulSoup
from num2words import num2words
from babel.dates import format_date
from babel.numbers import format_currency
from htmldocx import HtmlToDocx

from odoo import _, api, fields, models
from odoo.tools.safe_eval import safe_eval, time
from odoo.exceptions import ValidationError


class IrActionsReport(models.Model):
    _inherit = "ir.actions.report"

    report_type = fields.Selection(
        selection_add=[("docx", "DOCX")], ondelete={"docx": "cascade"}
    ) # add docx type
    report_docx_template = fields.Binary(string="Report DOCX Template")
    report_docx_template_name = fields.Char(string="Report DOCX Template Name")
    docx_merge_mode = fields.Selection(
        [("composer", "Composer"), ("zip", "Zip")],
        string="DOCX Mode",
        default="composer",
    )

    @api.constrains("report_type")
    def _check_report_type(self):
        for rec in self:
            if (
                rec.report_type == "docx"
                and not rec.report_docx_template
                and not rec.report_docx_template_name.endswith(".docx")
            ):
                raise ValidationError(_("Please upload a DOCX template."))

    def _render_docx(self, report_ref, docids, data):
        report = self._get_report_from_name(report_ref)
        template = report.report_docx_template

        if not template:
            raise ValueError("No DOCX template found.")

        doc_template = DocxTemplate(BytesIO(base64.b64decode(template)))
        doc_obj = self.env[report.model].browse(docids)
        render_image = partial(self._render_image, doc_template)
        html2docx = partial(self._render_html_as_subdoc, doc_template)

        context = {
            "spelled_out": self._spelled_out,
            "parsehtml": self._parse_html,
            "formatdate": self._formatdate,
            "company": self.env.company,
            "lang": self._context.get("lang", "es_ES"),
            "sysdate": fields.Datetime.now(),
            "render_image": render_image,
            "html2docx": html2docx,
            "convert_currency": self._convert_currency,
        }

        if report.docx_merge_mode == "composer":
            return self._render_composer_mode(doc_template, doc_obj, data, context)
        elif report.docx_merge_mode == "zip":
            return self._render_zip_mode(
                doc_template,
                doc_obj,
                data,
                context,
                report_name=report.print_report_name,
            )
        else:
            return self._render_docx_to_pdf_mode(doc_template, doc_obj, data, context)

    def _render_composer_mode(self, doc_template, doc_obj, data, context):
        master_doc = Document()
        composer = Composer(master_doc)

        for idx, obj in enumerate(doc_obj):
            context = {
                **context,
                "docs": obj,
                "data": data
            }

            temp = BytesIO()
            doc_template.render(context)
            doc_template.save(temp)
            temp.seek(0)

            if len(doc_obj) == 1:
                return temp.read()
            else:
                if idx == 0:
                    master_doc = Document(temp)
                    composer = Composer(master_doc)
                else:
                    doc_to_append = Document(temp)
                    master_doc.add_page_break()
                    composer.append(doc_to_append)

        temp_output = BytesIO()
        composer.save(temp_output)
        temp_output.seek(0)

        return temp_output.read()

    def _render_zip_mode(
        self, doc_template, doc_obj, data, context, report_name="report"
    ):
        docx_files = []

        for obj in doc_obj:
            context = {
                **context,
                "docs": obj,
                "data": data
            }

            temp = BytesIO()
            doc_template.render(context)
            doc_template.save(temp)
            temp.seek(0)
            docx_files.append(temp.read())

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            for idx, docx_file in enumerate(docx_files):
                name = safe_eval(report_name, {"object": doc_obj[idx], "time": time})
                filename = "%s.%s" % (name, "docx")
                zip_file.writestr(filename, docx_file)

        zip_buffer.seek(0)

        return zip_buffer.read()

    # Next Improvement
    def _render_docx_to_pdf_mode(self, doc_template, doc_obj, data, context):
        raise NotImplementedError("This method is not implemented yet.")

    def _convert_docx_to_pdf(self, docx_stream):
        raise NotImplementedError("This method is not implemented yet.")

    @staticmethod
    def _render_image(tpl, imgb64, width=None, height=None):
        width = Mm(width) if width else None
        height = Mm(height) if height else None
        
        if not imgb64:
            return ''

        image_stream = BytesIO(base64.b64decode(imgb64))
        return InlineImage(
            tpl, image_descriptor=image_stream, width=width, height=height
        )

    @staticmethod
    def _parse_html(html):
        if not html:
            return ""
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text()

    @staticmethod
    def _formatdate(date_required=fields.Datetime.today(), format="full", lang="es_ES"):
        return format_date(date_required, format=format, locale=lang)

    @staticmethod
    def _spelled_out(number, lang="es_ES", to="cardinal"):
        return num2words(number, lang=lang, to=to)
    
    @staticmethod
    def _convert_currency(number, currency_field, locale='es_ES'):
        return format_currency(number, currency_field.name, locale=locale)

    @staticmethod
    def _render_html_as_subdoc(tpl, html_code=None):
        if not (
            isinstance(html_code, str)
            and bool(BeautifulSoup(html_code, "html.parser").find())
        ):
            return ""

        temp = BytesIO()
        desc_document = Document()
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html_code, desc_document)
        desc_document.save(temp)
        temp.seek(0)
        return tpl.new_subdoc(temp)
